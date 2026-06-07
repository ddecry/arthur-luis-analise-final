from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
REPORT = ROOT / "relatorio"
ASSETS = REPORT / "assets"
OUT = REPORT / "arthur-luis-felipe-relatorio.docx"
DATA = json.loads((REPORT / "resumo_analise.json").read_text(encoding="utf-8"))

ACCENT = RGBColor(31, 78, 121)
GRAY = RGBColor(90, 90, 90)


def fix(text):
    return (
        str(text)
        .replace("Sęnior", "Sênior")
        .replace("Cięncia", "Ciência")
        .replace("Gestăo", "Gestão")
    )


def money(value):
    return f"R$ {float(value):,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def nint(value):
    return f"{int(value):,}".replace(",", ".")


def dec(value, digits=3):
    return f"{float(value):.{digits}f}".replace(".", ",")


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color or RGBColor(0, 0, 0)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph(doc, text="", first=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(7)
    if first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    set_run(r, 11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, 13 if level == 1 else 11.5, True, color=ACCENT)
    return p


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_border(cell)
            if row_index == 0:
                shade(cell, "F2F2F2")
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.05
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run(run, 9.5, row_index == 0)
    return table


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, 9.5, italic=True, color=GRAY)


def figure(doc, filename, cap, width_cm):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    caption(doc, cap)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.header.paragraphs[0].text = ""
section.footer.paragraphs[0].text = ""

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(7)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
run = title.add_run("State of Data Brazil 2024: fatores associados ao salário")
set_run(run, 15, True, color=ACCENT)

heading(doc, "Resumo")
paragraph(
    doc,
    "Este relatório analisa a base State of Data Brazil 2024-2025 com o objetivo de identificar fatores associados ao salário de profissionais de dados no Brasil. A amostra analisada contém 4.863 respostas com faixa salarial informada. Os resultados indicam que senioridade e experiência acumulada em dados são os fatores mais associados à remuneração.",
)

heading(doc, "1. Introdução")
paragraph(
    doc,
    "O mercado de dados possui diferenças salariais relacionadas ao nível profissional, ao cargo, à experiência acumulada e a outros fatores de contexto. A pergunta central deste trabalho é: quais fatores mais influenciam o salário de um profissional de dados no Brasil em 2024?",
)
paragraph(
    doc,
    "A análise é descritiva e observacional. Portanto, os resultados indicam associações entre variáveis, mas não permitem afirmar causalidade.",
)

heading(doc, "2. Metodologia")
paragraph(
    doc,
    f"A base original possui {nint(DATA['total_respostas'])} respostas. Após filtrar os registros com faixa salarial informada, permaneceram {nint(DATA['coorte_salario'])} observações. As faixas salariais foram convertidas em pontos médios. Para a faixa aberta superior, acima de R$ 40.001, foi utilizado o valor de R$ 45.000.",
)
paragraph(
    doc,
    "Foram selecionadas variáveis relacionadas a idade, gênero, raça/cor, região, escolaridade, situação de trabalho, cargo, senioridade, experiência e ferramentas utilizadas. Não foram imputadas respostas ausentes em variáveis sensíveis ou condicionais do questionário.",
)

heading(doc, "3. Resultados")
heading(doc, "3.1 Estatísticas gerais", 2)
stats = DATA["estatisticas"]
make_table(
    doc,
    ["Variável", "n", "Média", "Mediana", "Desvio padrão"],
    [
        ["Salário estimado", nint(stats["salario_estimado"]["n"]), money(stats["salario_estimado"]["media"]), money(stats["salario_estimado"]["mediana"]), money(stats["salario_estimado"]["desvio_padrao"])],
        ["Idade", nint(stats["idade"]["n"]), dec(stats["idade"]["media"], 1), dec(stats["idade"]["mediana"], 1), dec(stats["idade"]["desvio_padrao"], 1)],
        ["Experiência em dados", nint(stats["experiencia_dados_anos"]["n"]), dec(stats["experiencia_dados_anos"]["media"], 1), dec(stats["experiencia_dados_anos"]["mediana"], 1), dec(stats["experiencia_dados_anos"]["desvio_padrao"], 1)],
    ],
)
caption(doc, "Tabela 1 - Estatísticas gerais da amostra analisada.")
paragraph(
    doc,
    "A mediana salarial estimada da amostra é de aproximadamente R$ 10.000. A média é maior que a mediana, indicando assimetria à direita na distribuição dos salários.",
)
figure(doc, "01_distribuicao_salario.png", "Figura 1 - Distribuição do salário mensal estimado.", 12.6)

heading(doc, "3.2 Senioridade", 2)
make_table(
    doc,
    ["Senioridade", "n", "Mediana salarial", "Média salarial"],
    [[fix(item["senioridade"]), nint(item["count"]), money(item["median"]), money(item["mean"])] for item in DATA["senioridade"]],
)
caption(doc, "Tabela 2 - Salário estimado por senioridade.")
paragraph(
    doc,
    "A senioridade apresenta relação clara com o salário estimado. A mediana passa de cerca de R$ 3.500 no nível júnior para cerca de R$ 14.000 no nível sênior.",
)
figure(doc, "02_salario_senioridade.png", "Figura 2 - Salário mediano estimado por senioridade.", 11.6)

heading(doc, "3.3 Cargo", 2)
make_table(
    doc,
    ["Família de cargo", "n", "Mediana salarial"],
    [[fix(item["cargo_padronizado"]), nint(item["count"]), money(item["median"])] for item in DATA["cargo"]],
)
caption(doc, "Tabela 3 - Salário estimado por família de cargo.")
paragraph(
    doc,
    "As famílias de cargo também apresentam diferenças relevantes. Ciência de Dados e IA, Engenharia de Dados e Produto e Gestão aparecem com medianas mais altas na amostra.",
)
figure(doc, "03_salario_cargo.png", "Figura 3 - Salário mediano estimado por família de cargo.", 12.9)

heading(doc, "3.4 Correlações", 2)
corr = DATA["correlacoes_spearman"]
make_table(
    doc,
    ["Variável", "Correlação de Spearman com salário"],
    [
        ["Senioridade", dec(corr["senioridade_ordem"])],
        ["Experiência em dados", dec(corr["experiencia_dados_anos"])],
        ["Idade", dec(corr["idade"])],
        ["Escolaridade", dec(corr["escolaridade_ordem"])],
    ],
)
caption(doc, "Tabela 4 - Correlações de Spearman com o salário estimado.")
paragraph(
    doc,
    "A maior associação observada foi com senioridade, com coeficiente de 0,734. Em seguida aparece experiência em dados, com coeficiente de 0,658. Esses resultados reforçam que a maturidade profissional é o principal fator associado à remuneração na base analisada.",
)

heading(doc, "4. Discussão")
paragraph(
    doc,
    "Os resultados indicam que o salário de profissionais de dados está mais fortemente associado à senioridade e à experiência acumulada. Cargo e região ajudam a contextualizar diferenças, mas não devem ser interpretados isoladamente.",
)
paragraph(
    doc,
    "Variáveis como gênero e raça/cor exigem cuidado ético na interpretação. Como a base é observacional, diferenças nesses recortes podem estar relacionadas também a cargo, senioridade, setor, região e outras variáveis não controladas.",
)

heading(doc, "5. Limitações")
paragraph(
    doc,
    "A pesquisa é voluntária e pode não representar perfeitamente todo o mercado brasileiro de dados. Além disso, os salários foram informados em faixas, e a conversão para pontos médios introduz aproximações, especialmente na faixa aberta superior.",
)

heading(doc, "6. Conclusão")
paragraph(
    doc,
    "A análise aponta que senioridade e experiência acumulada em dados são os fatores mais associados ao salário estimado de profissionais de dados no Brasil em 2024. A conclusão deve ser entendida como descritiva: os dados mostram padrões de associação, mas não comprovam causalidade.",
)

heading(doc, "Referências")
paragraph(
    doc,
    "DATA HACKERS; BAIN & COMPANY. State of Data Brazil 2024-2025. Kaggle. Disponível em: https://www.kaggle.com/datasets/datahackers/state-of-data-brazil-20242025.",
    first=False,
)
paragraph(doc, "Bibliotecas utilizadas: pandas, NumPy, Matplotlib e seaborn.", first=False)

doc.core_properties.title = ""
doc.core_properties.subject = ""
doc.core_properties.author = ""
doc.core_properties.keywords = ""
doc.save(OUT)
print(OUT)
